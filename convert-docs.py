"""
将 docx 文件转换为 markdown 格式
"""
import os
from docx import Document
from pathlib import Path

def docx_to_markdown(docx_path: str) -> str:
    """将 docx 文件转换为 markdown 文本"""
    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 根据样式判断标题级别
        style = para.style.name if para.style else ""

        if "Heading 1" in style or "标题 1" in style:
            lines.append(f"# {text}")
        elif "Heading 2" in style or "标题 2" in style:
            lines.append(f"## {text}")
        elif "Heading 3" in style or "标题 3" in style:
            lines.append(f"### {text}")
        else:
            lines.append(text)

        lines.append("")  # 空行

    # 处理表格
    for table in doc.tables:
        # 表头
        if table.rows:
            header_row = table.rows[0]
            headers = [cell.text.strip() for cell in header_row.cells]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

            # 数据行
            for row in table.rows[1:]:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")

            lines.append("")

    return "\n".join(lines)

def convert_all_docx(input_dir: str, output_dir: str = None):
    """转换目录下所有 docx 文件"""
    input_path = Path(input_dir)

    if output_dir:
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
    else:
        output_path = input_path

    docx_files = list(input_path.glob("*.docx"))
    print(f"找到 {len(docx_files)} 个 docx 文件")

    for docx_file in docx_files:
        print(f"转换: {docx_file.name}")

        try:
            markdown_content = docx_to_markdown(str(docx_file))

            # 保存为 md 文件
            md_filename = docx_file.stem + ".md"
            md_path = output_path / md_filename

            with open(md_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)

            print(f"  -> {md_path}")
        except Exception as e:
            print(f"  错误: {e}")

    print("转换完成！")

if __name__ == "__main__":
    # 转换 my-designs 目录下的所有 docx
    convert_all_docx(
        "docs/my-designs",
        "docs/my-designs"  # 输出到同一目录
    )
